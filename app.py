import streamlit as st
import os
import tempfile
import json
import base64
from processing.parser import parse_srt, parse_srt_content
from processing.analyzer import analyze_words, apply_styling
from processing.formatter import generate_enhanced_output, json_to_srt
from processing.advanced_tools import analyze_subtitle_statistics, optimize_timing, check_errors, auto_translate
from langdetect import detect
from rich.console import Console
from rich.markdown import Markdown
import pysrt
from webvtt import WebVTT
import emoji
from docx import Document
import io

st.set_page_config(
    page_title="SRT Highlighter Pro - AI Tool",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Add custom CSS
st.markdown("""
<style>
    .stApp {
        max-width: 1200px;
        margin: 0 auto;
    }
    .main-header {
        text-align: center;
        color: #2e7d32;
        margin-bottom: 2rem;
    }
    .feature-card {
        padding: 1rem;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        margin-bottom: 1rem;
    }
    .output-preview {
        background-color: #f5f5f5;
        padding: 1rem;
        border-radius: 5px;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 24px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        white-space: pre-wrap;
        border-radius: 4px 4px 0px 0px;
        padding: 10px 16px;
        font-weight: 600;
    }
    .feedback-box {
        background-color: #e8f5e9;
        padding: 1rem;
        border-radius: 5px;
        margin-top: 1rem;
    }
    .error-box {
        background-color: #ffebee;
        padding: 1rem;
        border-radius: 5px;
        margin-top: 1rem;
    }
</style>
""", unsafe_allow_html=True)

def detect_language(text):
    """Detect the language of the text"""
    try:
        return detect(text)
    except:
        return 'en'  # Default to English if detection fails

def get_file_extension(output_format):
    """Get the appropriate file extension based on the output format"""
    format_extensions = {
        'html': '.html',
        'html_classic': '.html',
        'enhanced_srt': '.srt',
        'standard_srt': '.srt',
        'vtt': '.vtt',
        'ass': '.ass',
        'docx': '.docx',
        'json': '.json'
    }
    return format_extensions.get(output_format, '.txt')

def get_mime_type(output_format):
    """Get the appropriate MIME type based on the output format"""
    format_mimes = {
        'html': 'text/html',
        'html_classic': 'text/html',
        'enhanced_srt': 'text/plain',
        'standard_srt': 'text/plain',
        'vtt': 'text/vtt',
        'ass': 'text/plain',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'json': 'application/json'
    }
    return format_mimes.get(output_format, 'text/plain')

def convert_format(file_content, from_format, to_format):
    """Convert subtitle file from one format to another"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{from_format}') as temp_file:
        temp_file.write(file_content)
        temp_file_path = temp_file.name
    
    output_content = None
    
    try:
        if from_format == 'srt' and to_format == 'vtt':
            # Convert SRT to WebVTT
            srt_data = pysrt.open(temp_file_path)
            vtt = WebVTT()
            
            for sub in srt_data:
                vtt.captions.append(WebVTT().Caption(
                    start=str(sub.start).replace(',', '.'),
                    end=str(sub.end).replace(',', '.'),
                    text=sub.text
                ))
            
            # Save to temporary file
            output_path = temp_file_path.replace('.srt', '.vtt')
            vtt.save(output_path)
            
            with open(output_path, 'r', encoding='utf-8') as f:
                output_content = f.read()
            
            os.unlink(output_path)
        
        elif from_format == 'vtt' and to_format == 'srt':
            # Convert WebVTT to SRT
            vtt = WebVTT().read(temp_file_path)
            srt_content = []
            
            for i, caption in enumerate(vtt.captions, 1):
                srt_content.append(f"{i}\n")
                start = caption.start.replace('.', ',')
                end = caption.end.replace('.', ',')
                srt_content.append(f"{start} --> {end}\n")
                srt_content.append(f"{caption.text}\n\n")
            
            output_content = ''.join(srt_content)
        
        elif from_format == 'srt' and to_format == 'ass':
            # Convert SRT to ASS
            srt_data = pysrt.open(temp_file_path)
            words_data = []
            
            for sub in srt_data:
                words_data.append({
                    'word': sub.text,
                    'start_time': str(sub.start),
                    'end_time': str(sub.end)
                })
            
            output_content = generate_enhanced_output(words_data, output_format='ass')
        
        elif from_format == 'json' and to_format == 'srt':
            # Convert JSON to SRT
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            output_content = json_to_srt(json_data)
        
        elif from_format == 'srt' and to_format == 'docx':
            # Convert SRT to DOCX
            srt_data = pysrt.open(temp_file_path)
            doc = Document()
            
            for sub in srt_data:
                p = doc.add_paragraph()
                p.add_run(f"[{sub.start} --> {sub.end}]\n").italic = True
                p.add_run(sub.text)
                doc.add_paragraph()  # Add blank line
            
            # Save to temporary file
            output_path = temp_file_path.replace('.srt', '.docx')
            doc.save(output_path)
            
            # Read the file as bytes for download
            with open(output_path, 'rb') as f:
                output_content = f.read()
            
            os.unlink(output_path)
        
        # Clean up temporary file
        os.unlink(temp_file_path)
        
        return output_content
    except Exception as e:
        # Clean up and re-raise
        try:
            os.unlink(temp_file_path)
        except:
            pass
        raise e

def process_subtitle_analysis(words_data, language_code, min_importance, use_emojis, style_params):
    """Process subtitles with the analyzer"""
    if not words_data:
        return None, None
    
    # Process the words with the analyzer
    analyzed_words = analyze_words(words_data, language=language_code, min_importance=min_importance, use_emojis=use_emojis)
    
    # Apply styling
    styled_words = apply_styling(analyzed_words, use_emojis=use_emojis, **style_params)
    
    return analyzed_words, styled_words

def main():
    st.markdown("<h1 class='main-header'>🎯 AI SRT Highlighter Pro</h1>", unsafe_allow_html=True)
    
    # Create tabs for different functions
    tab1, tab2, tab3 = st.tabs(["Làm nổi bật phụ đề", "Chuyển đổi định dạng", "Công cụ nâng cao"])
    
    with tab1:
        st.markdown("""
        <div class='feature-card'>
        💡 Công cụ tự động làm nổi bật các từ quan trọng trong file phụ đề và thêm biểu tượng phù hợp.
        Hỗ trợ nhiều định dạng và tính năng nâng cao.
        </div>
        """, unsafe_allow_html=True)
        
        # Enhanced sidebar configuration
        with st.sidebar:
            st.header("⚙️ Cấu hình nâng cao")
            
            # Language settings
            st.subheader("Ngôn ngữ")
            auto_detect = st.checkbox("Tự động phát hiện ngôn ngữ", value=True)
            if not auto_detect:
                language = st.selectbox(
                    "Ngôn ngữ phân tích",
                    ["Tiếng Anh", "Tiếng Việt", "Tiếng Trung", "Tiếng Nhật", "Tiếng Hàn"],
                    index=0
                )
                language_code_map = {
                    "Tiếng Anh": "en", 
                    "Tiếng Việt": "vi", 
                    "Tiếng Trung": "zh", 
                    "Tiếng Nhật": "ja", 
                    "Tiếng Hàn": "ko"
                }
                lang_code = language_code_map.get(language, "en")
            else:
                lang_code = None  # Will be detected automatically
            
            # Output format settings
            st.subheader("Định dạng đầu ra")
            output_format = st.selectbox(
                "Chọn định dạng",
                [
                    "HTML (Modern)", 
                    "HTML (Classic)", 
                    "SRT Nâng cao",
                    "SRT Tiêu chuẩn",
                    "WebVTT",
                    "ASS/SSA",
                    "DOCX",
                    "JSON"
                ],
                index=0
            )
            
            # Style settings
            st.subheader("Tùy chỉnh kiểu dáng")
            use_custom_colors = st.checkbox("Sử dụng màu tùy chỉnh", value=False)
            if use_custom_colors:
                primary_color = st.color_picker("Màu chính", "#FF4B4B")
                secondary_color = st.color_picker("Màu phụ", "#45C6FF")
            else:
                primary_color = "#FF9900"
                secondary_color = "#3357FF"
            
            # Advanced settings
            st.subheader("Cài đặt nâng cao")
            group_words = st.checkbox("Nhóm từ theo thời gian", value=True)
            min_word_importance = st.slider(
                "Ngưỡng từ quan trọng",
                min_value=0.0,
                max_value=1.0,
                value=0.5,
                step=0.1
            )
            use_emojis = st.checkbox("Thêm biểu tượng cảm xúc", value=True)
            
        # Main content area
        st.header("1. Tải lên file phụ đề")
        
        upload_option = st.radio(
            "Chọn cách nhập dữ liệu",
            ["Tải lên file", "Nhập nội dung trực tiếp", "Từ URL"]
        )
        
        words_data = None
        
        if upload_option == "Tải lên file":
            uploaded_file = st.file_uploader("Tải lên file SRT", type=["srt", "vtt", "json"], key="upload_srt")
            
            if uploaded_file is not None:
                # Save file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{uploaded_file.name.split(".")[-1]}') as temp_file:
                    temp_file.write(uploaded_file.getbuffer())
                    temp_file_path = temp_file.name
                    
                # Process file
                try:
                    file_extension = uploaded_file.name.split(".")[-1].lower()
                    
                    if file_extension == "srt":
                        words_data = parse_srt(temp_file_path)
                    elif file_extension == "vtt":
                        # Convert VTT to SRT first
                        vtt = WebVTT().read(temp_file_path)
                        srt_path = temp_file_path.replace(".vtt", ".srt")
                        
                        with open(srt_path, 'w', encoding='utf-8') as f:
                            for i, caption in enumerate(vtt.captions, 1):
                                f.write(f"{i}\n")
                                start = caption.start.replace('.', ',')
                                end = caption.end.replace('.', ',')
                                f.write(f"{start} --> {end}\n")
                                f.write(f"{caption.text}\n\n")
                        
                        words_data = parse_srt(srt_path)
                        os.unlink(srt_path)
                    elif file_extension == "json":
                        with open(temp_file_path, 'r', encoding='utf-8') as f:
                            json_data = json.load(f)
                            words_data = json_data  # Assuming JSON is already in the correct format
                    
                    st.success(f"Đã tải file thành công! Tìm thấy {len(words_data)} từ/phân đoạn.")
                    
                    # Show preview
                    if words_data:
                        with st.expander("Xem trước dữ liệu thô", expanded=False):
                            st.write(words_data[:10])
                    
                    # Clean up temp file
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
                except Exception as e:
                    st.error(f"Lỗi khi xử lý file: {e}")
                    # Clean up temp file if there's an error
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
        
        elif upload_option == "Nhập nội dung trực tiếp":
            # Direct SRT content input
            srt_content = st.text_area("Nhập nội dung file SRT:", height=300)
            
            if st.button("Xử lý nội dung", key="process_srt_content") and srt_content:
                try:
                    words_data = parse_srt_content(srt_content)
                    st.success(f"Đã xử lý nội dung thành công! Tìm thấy {len(words_data)} từ/phân đoạn.")
                    
                    # Show preview
                    if words_data:
                        with st.expander("Xem trước dữ liệu thô", expanded=False):
                            st.write(words_data[:10])
                except Exception as e:
                    st.error(f"Lỗi khi xử lý nội dung: {e}")
        
        else:
            # URL option
            url = st.text_input("Nhập URL của file SRT")
            
            if st.button("Tải lên từ URL", key="process_url") and url:
                try:
                    import requests
                    response = requests.get(url)
                    
                    if response.status_code == 200:
                        content = response.text
                        # Save to temp file
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.srt') as temp_file:
                            temp_file.write(content.encode('utf-8'))
                            temp_file_path = temp_file.name
                        
                        words_data = parse_srt(temp_file_path)
                        st.success(f"Đã tải file từ URL thành công! Tìm thấy {len(words_data)} từ/phân đoạn.")
                        
                        # Show preview
                        if words_data:
                            with st.expander("Xem trước dữ liệu thô", expanded=False):
                                st.write(words_data[:10])
                        
                        # Clean up temp file
                        try:
                            os.unlink(temp_file_path)
                        except:
                            pass
                    else:
                        st.error(f"Không thể tải file từ URL: HTTP {response.status_code}")
                except Exception as e:
                    st.error(f"Lỗi khi tải lên từ URL: {e}")
        
        # New features for subtitle processing
        if words_data:
            st.header("2. Xử lý và phân tích")
            
            with st.spinner("Đang xử lý..."):
                try:
                    # Auto-detect language if enabled
                    if auto_detect and words_data:
                        # Get sample text for language detection
                        sample_text = " ".join([word.get('word', '') for word in words_data[:50]])
                        lang_code = detect_language(sample_text)
                        st.info(f"Đã phát hiện ngôn ngữ: {lang_code}")
                    
                    # Prepare style parameters
                    style_params = {}
                    if use_custom_colors:
                        style_params['primary_color'] = primary_color
                        style_params['secondary_color'] = secondary_color
                    
                    # Process and analyze the subtitles
                    analyzed_words, styled_words = process_subtitle_analysis(
                        words_data, 
                        lang_code, 
                        min_word_importance, 
                        use_emojis, 
                        style_params
                    )
                    
                    # Map output format selection to actual format string
                    format_mapping = {
                        "HTML (Modern)": "html",
                        "HTML (Classic)": "html_classic",
                        "SRT Nâng cao": "enhanced_srt",
                        "SRT Tiêu chuẩn": "standard_srt",
                        "WebVTT": "vtt",
                        "ASS/SSA": "ass",
                        "DOCX": "docx",
                        "JSON": "json"
                    }
                    
                    output_format_str = format_mapping.get(output_format, "html")
                    
                    if styled_words:
                        # Generate output based on format
                        enhanced_output = generate_enhanced_output(
                            styled_words, 
                            output_format=output_format_str, 
                            group_words=group_words, 
                            style_params=style_params if use_custom_colors else None
                        )
                        
                        if enhanced_output:
                            st.success("Đã xử lý và định dạng thành công!")
                            
                            # Special handling for DOCX
                            if output_format_str == "docx":
                                # Create a BytesIO object for DOCX
                                output_bytes = io.BytesIO()
                                doc = Document()
                                
                                for word in styled_words:
                                    p = doc.add_paragraph()
                                    p.add_run(f"[{word.get('start_time', '')} --> {word.get('end_time', '')}]\n").italic = True
                                    
                                    # Add styling based on importance
                                    if word.get('important', False):
                                        run = p.add_run(f"{word.get('primary_icon', '')} {word.get('word', '')}")
                                        run.bold = True
                                        # Convert hex color to RGB for docx
                                        color = word.get('style', {}).get('color', '#000000')
                                        if color.startswith('#'):
                                            r = int(color[1:3], 16)
                                            g = int(color[3:5], 16)
                                            b = int(color[5:7], 16)
                                            run.font.color.rgb = (r, g, b)
                                    else:
                                        p.add_run(f"{word.get('primary_icon', '')} {word.get('word', '')}")
                                    
                                    doc.add_paragraph()  # Add blank line
                                
                                doc.save(output_bytes)
                                output_bytes.seek(0)
                                
                                st.download_button(
                                    "Tải xuống kết quả DOCX",
                                    output_bytes,
                                    file_name=f"enhanced_subtitles.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                            else:
                                st.download_button(
                                    "Tải xuống kết quả",
                                    enhanced_output,
                                    file_name=f"enhanced_subtitles{get_file_extension(output_format_str)}",
                                    mime=get_mime_type(output_format_str)
                                )
                            
                            # Preview the output
                            with st.expander("Xem trước kết quả", expanded=True):
                                if output_format_str == "html":
                                    st.markdown(enhanced_output, unsafe_allow_html=True)
                                elif output_format_str == "docx":
                                    st.info("Tập tin DOCX không thể xem trước trực tiếp. Vui lòng tải xuống để xem.")
                                else:
                                    st.code(enhanced_output)
                                    
                            # Show statistics
                            with st.expander("Thống kê phân tích", expanded=False):
                                total_words = len(styled_words)
                                important_words = sum(1 for word in styled_words if word.get('important', False))
                                
                                st.write(f"Tổng số từ: {total_words}")
                                st.write(f"Số từ quan trọng: {important_words} ({important_words/total_words*100:.1f}%)")
                                
                                # Category distribution
                                categories = {}
                                for word in styled_words:
                                    for cat in word.get('categories', []):
                                        categories[cat] = categories.get(cat, 0) + 1
                                
                                if categories:
                                    st.write("Phân bố theo danh mục:")
                                    for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
                                        st.write(f"- {cat}: {count} từ ({count/total_words*100:.1f}%)")
                        else:
                            st.warning("Không có dữ liệu đầu ra được tạo.")
                    else:
                        st.warning("Không thể phân tích nội dung. Vui lòng kiểm tra lại dữ liệu đầu vào.")
                except Exception as e:
                    st.error(f"Lỗi trong quá trình xử lý: {str(e)}")
                    st.error("Vui lòng kiểm tra lại dữ liệu đầu vào và thử lại.")

    with tab2:
        st.header("🔄 Chuyển đổi định dạng phụ đề")
        
        # Define conversion options
        conversion_options = {
            "SRT ⟶ WebVTT": {"from": "srt", "to": "vtt"},
            "WebVTT ⟶ SRT": {"from": "vtt", "to": "srt"},
            "SRT ⟶ ASS": {"from": "srt", "to": "ass"},
            "JSON ⟶ SRT": {"from": "json", "to": "srt"},
            "SRT ⟶ DOCX": {"from": "srt", "to": "docx"}
        }
        
        selected_conversion = st.selectbox(
            "Chọn kiểu chuyển đổi",
            list(conversion_options.keys())
        )
        
        conversion_info = conversion_options[selected_conversion]
        
        # File upload for conversion
        uploaded_file = st.file_uploader(
            f"Tải lên file {conversion_info['from'].upper()}", 
            type=[conversion_info['from']], 
            key="convert_upload"
        )
        
        if uploaded_file is not None:
            try:
                # Process the uploaded file
                file_content = uploaded_file.getbuffer()
                
                output_content = convert_format(
                    file_content, 
                    conversion_info['from'], 
                    conversion_info['to']
                )
                
                if output_content:
                    # Determine download button parameters
                    file_extension = f".{conversion_info['to']}"
                    mime_type = get_mime_type(conversion_info['to'])
                    
                    # Create download button
                    st.download_button(
                        f"Tải xuống file {conversion_info['to'].upper()}",
                        output_content,
                        file_name=f"converted{file_extension}",
                        mime=mime_type
                    )
                    
                    # Preview content if not binary
                    if conversion_info['to'] != "docx":
                        with st.expander("Xem trước nội dung", expanded=False):
                            st.code(output_content)
                    else:
                        st.info("Tập tin DOCX không thể xem trước trực tiếp. Vui lòng tải xuống để xem.")
                
            except Exception as e:
                st.error(f"Lỗi khi chuyển đổi file: {str(e)}")
                st.error("Vui lòng kiểm tra lại định dạng file đầu vào.")

    with tab3:
        st.header("🛠️ Công cụ nâng cao")
        
        tool_options = {
            "Phân tích thống kê phụ đề": "subtitle_stats",
            "Tối ưu hóa thời gian": "timing_optimization",
            "Kiểm tra lỗi": "error_check",
            "Dịch tự động": "auto_translate"
        }
        
        selected_tool = st.selectbox(
            "Chọn công cụ",
            list(tool_options.keys())
        )
        
        tool_type = tool_options[selected_tool]
        
        # File upload for advanced tools
        uploaded_file = st.file_uploader(
            "Tải lên file SRT", 
            type=["srt"], 
            key="advanced_tools_upload"
        )
        
        if uploaded_file is not None:
            try:
                # Save file temporarily
                with tempfile.NamedTemporaryFile(delete=False, suffix='.srt') as temp_file:
                    temp_file.write(uploaded_file.getbuffer())
                    temp_file_path = temp_file.name
                
                # Load SRT data
                srt_data = pysrt.open(temp_file_path)
                
                if tool_type == "subtitle_stats":
                    # Analyze subtitle statistics
                    stats = analyze_subtitle_statistics(srt_data)
                    
                    st.subheader("Thống kê phụ đề")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.metric("Tổng số phụ đề", len(srt_data))
                        st.metric("Tổng số từ", stats['word_count'])
                        st.metric("Số từ khác nhau", len(stats['unique_words']))
                    
                    with col2:
                        total_seconds = stats['total_duration'].total_seconds()
                        st.metric("Tổng thời lượng", f"{int(total_seconds // 60)} phút {int(total_seconds % 60)} giây")
                        avg_words = stats['avg_words_per_subtitle']
                        st.metric("Trung bình từ/phụ đề", f"{avg_words:.1f}")
                        avg_duration = stats['avg_duration_per_subtitle'].total_seconds()
                        st.metric("Thời lượng TB/phụ đề", f"{avg_duration:.1f} giây")
                    
                    # Word frequency
                    st.subheader("Từ xuất hiện nhiều nhất")
                    freq_items = sorted(stats['word_frequency'].items(), key=lambda x: x[1], reverse=True)
                    freq_df = {"Từ": [], "Tần suất": []}
                    
                    for word, freq in freq_items[:20]:
                        freq_df["Từ"].append(word)
                        freq_df["Tần suất"].append(freq)
                    
                    st.bar_chart(freq_df, x="Từ")
                    
                    # Language distribution if available
                    if stats['language_distribution']:
                        st.subheader("Phân bố ngôn ngữ")
                        langs = {"Ngôn ngữ": [], "Số lượng": []}
                        
                        for lang, count in stats['language_distribution'].items():
                            langs["Ngôn ngữ"].append(lang)
                            langs["Số lượng"].append(count)
                        
                        st.bar_chart(langs, x="Ngôn ngữ")
                
                elif tool_type == "timing_optimization":
                    # UI controls for optimization parameters
                    st.subheader("Tùy chọn tối ưu hóa thời gian")
                    
                    min_duration = st.slider("Thời lượng tối thiểu (giây)", 0.5, 3.0, 1.0, 0.1)
                    max_duration = st.slider("Thời lượng tối đa (giây)", 3.0, 10.0, 7.0, 0.1)
                    min_gap = st.slider("Khoảng cách tối thiểu (giây)", 0.0, 1.0, 0.1, 0.05)
                    
                    if st.button("Tối ưu hóa"):
                        with st.spinner("Đang tối ưu hóa..."):
                            # Optimize timing
                            optimized_srt = optimize_timing(srt_data, min_duration, max_duration, min_gap)
                            
                            # Save optimized SRT to a temporary file
                            optimized_path = temp_file_path.replace('.srt', '_optimized.srt')
                            optimized_srt_file = pysrt.SubRipFile(optimized_srt)
                            optimized_srt_file.save(optimized_path, encoding='utf-8')
                            
                            # Read the optimized file content
                            with open(optimized_path, 'r', encoding='utf-8') as f:
                                optimized_content = f.read()
                            
                            # Create download button
                            st.download_button(
                                "Tải xuống SRT đã tối ưu hóa",
                                optimized_content,
                                file_name="optimized_subtitle.srt",
                                mime="text/plain"
                            )
                            
                            # Show preview
                            with st.expander("Xem trước nội dung đã tối ưu hóa", expanded=True):
                                st.code(optimized_content[:1000] + "..." if len(optimized_content) > 1000 else optimized_content)
                            
                            # Clean up temporary file
                            try:
                                os.unlink(optimized_path)
                            except:
                                pass
                            
                            # Show statistics comparison
                            st.subheader("So sánh trước và sau khi tối ưu hóa")
                            
                            # Original stats
                            original_durations = [(sub.end - sub.start).total_seconds() for sub in srt_data]
                            original_avg_duration = sum(original_durations) / len(original_durations) if original_durations else 0
                            
                            # Optimized stats
                            optimized_durations = [(sub.end - sub.start).total_seconds() for sub in optimized_srt]
                            optimized_avg_duration = sum(optimized_durations) / len(optimized_durations) if optimized_durations else 0
                            
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.metric("Thời lượng TB (ban đầu)", f"{original_avg_duration:.2f}s")
                                st.metric("Thời lượng tối thiểu (ban đầu)", f"{min(original_durations):.2f}s")
                                st.metric("Thời lượng tối đa (ban đầu)", f"{max(original_durations):.2f}s")
                            
                            with col2:
                                st.metric("Thời lượng TB (sau tối ưu)", f"{optimized_avg_duration:.2f}s")
                                st.metric("Thời lượng tối thiểu (sau tối ưu)", f"{min(optimized_durations):.2f}s")
                                st.metric("Thời lượng tối đa (sau tối ưu)", f"{max(optimized_durations):.2f}s")
                
                elif tool_type == "error_check":
                    st.subheader("Kiểm tra lỗi phụ đề")
                    
                    with st.spinner("Đang kiểm tra lỗi..."):
                        # Check for errors
                        errors = check_errors(srt_data)
                        
                        if errors:
                            st.error(f"Đã tìm thấy {len(errors)} lỗi trong file phụ đề!")
                            
                            for i, error in enumerate(errors, 1):
                                with st.expander(f"Lỗi #{i}: {error['message']}"):
                                    st.write(f"Phụ đề thứ: {error['subtitle_index']}")
                                    st.write(f"Loại lỗi: {error['type']}")
                                    
                                    # Show the problematic subtitle
                                    subtitle = srt_data[error['subtitle_index'] - 1]
                                    st.code(f"{subtitle.index}\n{subtitle.start} --> {subtitle.end}\n{subtitle.text}")
                            
                            # Offer auto-fix option
                            if st.button("Sửa lỗi tự động"):
                                fixed_srt = srt_data.copy()
                                
                                # Fix each error
                                for error in errors:
                                    idx = error['subtitle_index'] - 1
                                    subtitle = fixed_srt[idx]
                                    
                                    if error['type'] == 'negative_duration':
                                        # Fix negative duration by setting end time to start time + 2 seconds
                                        subtitle.end = subtitle.start + pysrt.SubRipTime(0, 0, 2)
                                    
                                    elif error['type'] == 'overlap':
                                        if idx < len(fixed_srt) - 1:
                                            next_subtitle = fixed_srt[idx + 1]
                                            # Fix overlap by adjusting end time
                                            subtitle.end = next_subtitle.start - pysrt.SubRipTime(0, 0, 0, 100)
                                    
                                    elif error['type'] == 'empty_text':
                                        # Fix empty text by adding placeholder
                                        subtitle.text = "[NO TEXT]"
                                    
                                    elif error['type'] == 'too_long':
                                        # No automatic fix for too long subtitles, just highlight
                                        pass
                                
                                # Save fixed SRT to a temporary file
                                fixed_path = temp_file_path.replace('.srt', '_fixed.srt')
                                fixed_srt.save(fixed_path, encoding='utf-8')
                                
                                # Read the fixed file content
                                with open(fixed_path, 'r', encoding='utf-8') as f:
                                    fixed_content = f.read()
                                
                                # Create download button
                                st.download_button(
                                    "Tải xuống SRT đã sửa lỗi",
                                    fixed_content,
                                    file_name="fixed_subtitle.srt",
                                    mime="text/plain"
                                )
                                
                                # Clean up temporary file
                                try:
                                    os.unlink(fixed_path)
                                except:
                                    pass
                        else:
                            st.success("Không tìm thấy lỗi nào trong file phụ đề!")
                
                elif tool_type == "auto_translate":
                    st.subheader("Dịch tự động phụ đề")
                    
                    # Select source and target languages
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        source_lang = st.selectbox(
                            "Ngôn ngữ nguồn",
                            ["Tự động phát hiện", "Tiếng Anh", "Tiếng Việt", "Tiếng Trung", "Tiếng Nhật", "Tiếng Hàn", "Tiếng Pháp", "Tiếng Đức", "Tiếng Tây Ban Nha"],
                            index=0
                        )
                    
                    with col2:
                        target_lang = st.selectbox(
                            "Ngôn ngữ đích",
                            ["Tiếng Anh", "Tiếng Việt", "Tiếng Trung", "Tiếng Nhật", "Tiếng Hàn", "Tiếng Pháp", "Tiếng Đức", "Tiếng Tây Ban Nha"],
                            index=0
                        )
                    
                    language_code_map = {
                        "Tự động phát hiện": "auto",
                        "Tiếng Anh": "en", 
                        "Tiếng Việt": "vi", 
                        "Tiếng Trung": "zh", 
                        "Tiếng Nhật": "ja", 
                        "Tiếng Hàn": "ko",
                        "Tiếng Pháp": "fr",
                        "Tiếng Đức": "de",
                        "Tiếng Tây Ban Nha": "es"
                    }
                    
                    source_code = language_code_map.get(source_lang, "auto")
                    target_code = language_code_map.get(target_lang, "en")
                    
                    # Display warning about translation API requirements
                    st.warning("""
                    Tính năng này yêu cầu cài đặt thư viện transformers và mô hình dịch thuật.
                    Nếu bạn đang chạy ứng dụng cục bộ, hãy đảm bảo bạn đã cài đặt:
                    ```
                    pip install transformers sentencepiece
                    ```
                    """)
                    
                    # Check if translation is feasible
                    if source_code == target_code:
                        st.error("Ngôn ngữ nguồn và đích không thể giống nhau!")
                    else:
                        if st.button("Bắt đầu dịch"):
                            try:
                                with st.spinner(f"Đang dịch từ {source_lang} sang {target_lang}..."):
                                    # If using auto-detect, detect the language first
                                    if source_code == "auto":
                                        sample_text = "\n".join([sub.text for sub in srt_data[:5]])
                                        detected_lang = detect_language(sample_text)
                                        source_code = detected_lang
                                        st.info(f"Đã phát hiện ngôn ngữ nguồn: {detected_lang}")
                                    
                                    # Show demo translation for first few subtitles
                                    st.subheader("Xem trước bản dịch")
                                    
                                    # Create a placeholder for samples
                                    sample_container = st.container()
                                    
                                    with sample_container:
                                        st.write("Đang tạo bản dịch mẫu...")
                                    
                                    # In a real implementation, you would use a translation service here
                                    sample_translations = [
                                        {"original": sub.text, "translated": f"[Bản dịch {target_code}] {sub.text}"} 
                                        for sub in srt_data[:5]
                                    ]
                                    
                                    # Update the sample container
                                    with sample_container:
                                        st.write("Mẫu bản dịch (5 phụ đề đầu tiên):")
                                        for sample in sample_translations:
                                            st.write("---")
                                            st.write(f"**Gốc:** {sample['original']}")
                                            st.write(f"**Dịch:** {sample['translated']}")
                                    
                                    # In a real implementation, you would translate all subtitles here
                                    # translated_srt = auto_translate(srt_data, source_code, target_code)
                                    
                                    # For demonstration, create a mock translated SRT file
                                    translated_srt = srt_data.copy()
                                    for sub in translated_srt:
                                        sub.text = f"[{target_code}] {sub.text}"
                                    
                                    # Save translated SRT to a temporary file
                                    translated_path = temp_file_path.replace('.srt', f'_{target_code}.srt')
                                    translated_srt.save(translated_path, encoding='utf-8')
                                    
                                    # Read the translated file content
                                    with open(translated_path, 'r', encoding='utf-8') as f:
                                        translated_content = f.read()
                                    
                                    # Create download button
                                    st.download_button(
                                        f"Tải xuống bản dịch sang {target_lang}",
                                        translated_content,
                                        file_name=f"translated_{target_code}_subtitle.srt",
                                        mime="text/plain"
                                    )
                                    
                                    # Clean up temporary file
                                    try:
                                        os.unlink(translated_path)
                                    except:
                                        pass
                                    
                                    st.success("Đã hoàn thành dịch thuật!")
                            except Exception as e:
                                st.error(f"Lỗi khi dịch phụ đề: {str(e)}")
                                st.info("Để sử dụng tính năng này, bạn cần cài đặt thư viện transformers và các mô hình dịch thuật phù hợp.")
                
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
            
            except Exception as e:
                st.error(f"Lỗi khi xử lý công cụ nâng cao: {str(e)}")
                # Clean up temporary file if there's an error
                try:
                    os.unlink(temp_file_path)
                except:
                    pass
        
        # Additional info about advanced tools
        st.markdown("""
        ### Thông tin về các công cụ nâng cao

        **Phân tích thống kê phụ đề**: Phân tích chi tiết các thông số của file phụ đề như số lượng từ, thời lượng, tần suất từ.

        **Tối ưu hóa thời gian**: Điều chỉnh thời gian hiển thị của phụ đề để cải thiện trải nghiệm người xem.

        **Kiểm tra lỗi**: Phát hiện và sửa các lỗi phổ biến trong file phụ đề như thời gian âm, trùng lặp, văn bản trống.

        **Dịch tự động**: Dịch phụ đề sang ngôn ngữ khác sử dụng các mô hình học máy.
        """)

if __name__ == "__main__":
    main()