import pysrt
from webvtt import WebVTT
import json
from docx import Document
from docx.shared import RGBColor
import re
from datetime import timedelta
from langdetect import detect
from transformers import pipeline

def analyze_subtitle_statistics(srt_data):
    """
    Analyze subtitle statistics including timing, word frequency, and complexity.
    """
    stats = {
        'total_duration': timedelta(),
        'word_count': 0,
        'unique_words': set(),
        'avg_words_per_subtitle': 0,
        'avg_duration_per_subtitle': timedelta(),
        'word_frequency': {},
        'language_distribution': {}
    }
    
    for subtitle in srt_data:
        duration = subtitle.end - subtitle.start
        stats['total_duration'] += duration
        
        words = re.findall(r'\w+', subtitle.text.lower())
        stats['word_count'] += len(words)
        stats['unique_words'].update(words)
        
        for word in words:
            stats['word_frequency'][word] = stats['word_frequency'].get(word, 0) + 1
            
        try:
            lang = detect(subtitle.text)
            stats['language_distribution'][lang] = stats['language_distribution'].get(lang, 0) + 1
        except:
            pass
    
    if len(srt_data) > 0:
        stats['avg_words_per_subtitle'] = stats['word_count'] / len(srt_data)
        stats['avg_duration_per_subtitle'] = stats['total_duration'] / len(srt_data)
    
    return stats

def optimize_timing(srt_data, min_duration=1.0, max_duration=7.0, min_gap=0.1):
    """
    Optimize subtitle timing for better readability.
    """
    optimized = []
    
    for i, subtitle in enumerate(srt_data):
        duration = subtitle.end - subtitle.start
        
        # Ensure minimum duration
        if duration.total_seconds() < min_duration:
            subtitle.end = subtitle.start + timedelta(seconds=min_duration)
        
        # Ensure maximum duration
        if duration.total_seconds() > max_duration:
            subtitle.end = subtitle.start + timedelta(seconds=max_duration)
        
        # Check gap with next subtitle
        if i < len(srt_data) - 1:
            next_subtitle = srt_data[i + 1]
            gap = next_subtitle.start - subtitle.end
            
            if gap.total_seconds() < min_gap:
                subtitle.end = next_subtitle.start - timedelta(seconds=min_gap)
        
        optimized.append(subtitle)
    
    return optimized

def check_errors(srt_data):
    """
    Check for common subtitle errors.
    """
    errors = []
    
    for i, subtitle in enumerate(srt_data):
        # Check duration
        duration = subtitle.end - subtitle.start
        if duration.total_seconds() < 0:
            errors.append({
                'type': 'negative_duration',
                'subtitle_index': i + 1,
                'message': f'Negative duration at subtitle {i + 1}'
            })
        
        # Check overlap with next subtitle
        if i < len(srt_data) - 1:
            next_subtitle = srt_data[i + 1]
            if subtitle.end > next_subtitle.start:
                errors.append({
                    'type': 'overlap',
                    'subtitle_index': i + 1,
                    'message': f'Overlap between subtitles {i + 1} and {i + 2}'
                })
        
        # Check for empty text
        if not subtitle.text.strip():
            errors.append({
                'type': 'empty_text',
                'subtitle_index': i + 1,
                'message': f'Empty text at subtitle {i + 1}'
            })
        
        # Check for excessive length
        words = len(subtitle.text.split())
        if words > 15:  # Generally accepted maximum words per subtitle
            errors.append({
                'type': 'too_long',
                'subtitle_index': i + 1,
                'message': f'Too many words ({words}) at subtitle {i + 1}'
            })
    
    return errors

def auto_translate(srt_data, target_lang='en'):
    """
    Automatically translate subtitles using transformers.
    """
    translator = pipeline("translation", model=f"Helsinki-NLP/opus-mt-{detect(srt_data[0].text)}-{target_lang}")
    
    translated = []
    for subtitle in srt_data:
        translation = translator(subtitle.text)[0]['translation_text']
        subtitle.text = translation
        translated.append(subtitle)
    
    return translated

def convert_format(input_data, from_format, to_format, **kwargs):
    """
    Convert between different subtitle formats.
    """
    supported_formats = ['srt', 'vtt', 'ass', 'json', 'docx']
    
    if from_format not in supported_formats or to_format not in supported_formats:
        raise ValueError(f"Unsupported format. Supported formats: {supported_formats}")
    
    # Implement conversion logic for each format pair
    if from_format == 'srt' and to_format == 'vtt':
        return srt_to_vtt(input_data)
    elif from_format == 'vtt' and to_format == 'srt':
        return vtt_to_srt(input_data)
    elif from_format == 'srt' and to_format == 'docx':
        return srt_to_docx(input_data, **kwargs)
    # Add more conversion pairs as needed
    
    raise NotImplementedError(f"Conversion from {from_format} to {to_format} not implemented yet")

def srt_to_vtt(srt_data):
    """Convert SRT to WebVTT format."""
    vtt = WebVTT()
    for subtitle in srt_data:
        vtt.captions.append(WebVTT.Caption(
            subtitle.start,
            subtitle.end,
            subtitle.text
        ))
    return vtt

def vtt_to_srt(vtt_data):
    """Convert WebVTT to SRT format."""
    srt = pysrt.SubRipFile()
    for i, caption in enumerate(vtt_data.captions, 1):
        srt.append(pysrt.SubRipItem(
            index=i,
            start=caption.start,
            end=caption.end,
            text=caption.text
        ))
    return srt

def srt_to_docx(srt_data, include_timestamps=True, highlight_important=True):
    """Convert SRT to DOCX format with formatting."""
    doc = Document()
    
    for subtitle in srt_data:
        paragraph = doc.add_paragraph()
        
        if include_timestamps:
            paragraph.add_run(f"[{subtitle.start} --> {subtitle.end}]\n").italic = True
        
        if highlight_important and hasattr(subtitle, 'important') and subtitle.important:
            run = paragraph.add_run(subtitle.text)
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red for important text
        else:
            paragraph.add_run(subtitle.text)
        
        doc.add_paragraph()  # Add blank line between subtitles
    
    return doc 